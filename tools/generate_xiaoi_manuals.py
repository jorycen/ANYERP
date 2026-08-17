from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
GREEN, DARK, NAVY, GRAY = "1AAD19", "116B1E", "17324D", "5B6670"
LIGHT, BLUEFILL, GOLD = "EEF8F0", "E8EEF5", "FFF8E8"
WIDTH = 9360


def font(run, size=11, color="222222", bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def cell_fill(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
    if shd.getparent() is None:
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell, color="D7DEE5", size="6"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def cell_margins(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = mar.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
        if node.getparent() is None:
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    pr = table._tbl.tblPr
    tw = pr.find(qn("w:tblW"))
    if tw is None:
        tw = OxmlElement("w:tblW")
    if tw.getparent() is None:
        pr.append(tw)
    tw.set(qn("w:w"), str(sum(widths)))
    tw.set(qn("w:type"), "dxa")
    ind = pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
    if ind.getparent() is None:
        pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
            if tcw.getparent() is None:
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[i]))
            tcw.set(qn("w:type"), "dxa")
            cell_margins(cell)
            cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    trpr.append(node)


def page_number(p):
    r = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    r._r.append(begin); r._r.append(instr); r._r.append(end)
    font(r, 9, GRAY)


def setup(doc, title):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, color, before, after in [(1, 16, DARK, 18, 10), (2, 13, DARK, 14, 7), (3, 12, NAVY, 10, 5)]:
        st = styles[f"Heading {level}"]
        st.font.name = "Aptos Display"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for sec in doc.sections:
        sec.page_width = Inches(8.5); sec.page_height = Inches(11)
        sec.top_margin = Inches(1); sec.bottom_margin = Inches(1); sec.left_margin = Inches(1); sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.492); sec.footer_distance = Inches(0.492)
        hp = sec.header.paragraphs[0]
        r = hp.add_run(title + "  ·  小艾 2.0")
        font(r, 9, GRAY, True)
        fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = fp.add_run("内部使用  ·  第 "); font(r, 9, GRAY); page_number(fp); r = fp.add_run(" 页"); font(r, 9, GRAY)
    doc.core_properties.title = title
    doc.core_properties.author = "小艾产品团队"


def para(doc, text="", size=11, color="222222", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.25
    if text:
        r = p.add_run(text); font(r, size, color, bold, italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}"); p.paragraph_format.keep_with_next = True
    r = p.add_run(text); font(r, {1:16,2:13,3:12}[level], {1:DARK,2:DARK,3:NAVY}[level], True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375); p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text); font(r, 10.5)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375); p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text); font(r, 10.5)
    return p


def callout(doc, title, text, fill=LIGHT, accent=GREEN):
    t = doc.add_table(rows=1, cols=1); geometry(t, [WIDTH]); c = t.cell(0, 0)
    cell_fill(c, fill); cell_border(c, accent, "12")
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.2
    r = p.add_run(title + "  "); font(r, 10.5, accent, True); r = p.add_run(text); font(r, 10.5, "28323B")
    para(doc, "", after=2)


def table(doc, headers, rows, widths, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers)); geometry(t, widths); repeat_header(t.rows[0])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; cell_fill(c, BLUEFILL); p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h); font(r, size, NAVY, True)
    for row in rows:
        cs = t.add_row().cells
        for i, value in enumerate(row):
            p = cs[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value)); font(r, size)
    para(doc, "", after=2)
    return t


def cover(doc, subtitle, audience):
    para(doc, "小艾 2.0", 31, NAVY, True, align=WD_ALIGN_PARAGRAPH.CENTER, before=64, after=8)
    para(doc, subtitle, 15, DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=25)
    para(doc, "让门店业务更快、更准、更容易协同", 11, GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=78)
    para(doc, "版本：2.0  |  文档日期：2026 年 7 月 28 日", 10.5, NAVY, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    para(doc, audience, 9.5, GRAY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=25)
    callout(doc, "说明", "本文依据当前小艾小程序已实现的功能整理。具体页面、数据范围和按钮权限以实际账号与部署版本为准。", LIGHT, GREEN)
    doc.add_page_break()


def build_product():
    d = Document(); setup(d, "产品说明书"); cover(d, "门店销售与经营协同助手 · 产品说明书", "适用对象：经销商、店长、店员及门店运营管理人员")
    heading(d, "目录")
    for x in ["一、产品定位", "二、核心功能", "三、销售与客户管理", "四、库存与供应链", "五、审批与经营分析", "六、采购、费用与基础管理", "七、权限、设备与注意事项"]: bullet(d, x)
    heading(d, "一、产品定位")
    para(d, "小艾 2.0 面向门店销售与经营管理场景，围绕“销售成交—库存流转—采购补给—审批留痕—经营复盘”建立统一入口。它既支持一线员工现场快速开单，也支持经销商和店长查看经营数据、处理审批和维护基础资料。")
    table(d, ["能力域", "主要功能", "价值"], [
        ("销售开单", "新建订单、客户来源、PN/SN、收款、定金、补贴、开票", "减少重复录入，确保订单完整"),
        ("订单运营", "订单查询、详情编辑、导出任务、图片与开票资料", "快速定位订单，支持后续补资料"),
        ("库存供应", "库存查询、采购入库、门店调拨、出库/收货凭证", "库存可查、可调、可追踪"),
        ("审批协同", "销售、采购、报销、商品、退单、资源、毛利调整", "关键动作有流程、有留痕"),
        ("经营分析", "经营报表、门店排行、品类/重点产品、员工销售、业绩上报", "支持日常复盘与管理决策"),
        ("基础管理", "门店、员工、用户、商品、客户来源、收款方式、补录项目", "统一业务口径，降低管理成本"),
    ], [1500, 4100, 3760])
    heading(d, "二、核心功能")
    for x in [
        "角色化工作台：首页按账号角色呈现订单、库存、调拨、入库、审批和报表入口。",
        "现场高效录入：支持商品搜索、PN/SN 扫码或 OCR 识别，适合门店现场操作。",
        "销售与库存联动：订单、入库、调拨和库存查询围绕商品编码建立关联。",
        "流程化经营：采购、报销、商品和销售等事项统一进入审批中心处理。",
        "数据化复盘：通过经营报表、业绩上报和商场销量查询查看经营结果。",
        "可配置基础资料：门店、商品、客户来源、收款方式和补录项目均可维护。",
    ]: bullet(d, x)
    heading(d, "三、销售与客户管理")
    heading(d, "3.1 新建订单", 2)
    para(d, "用于录入客户、商品、收款和补贴信息并形成销售订单。支持客户来源、会员称呼/ID、商品名称、PN、SN、IMEI、数量、单价、收款方式、定金抵扣、国补、教育补贴、开票和照片资料。")
    table(d, ["模块", "特点"], [
        ("商品信息", "支持搜索、扫码和多商品录入；SN 商品可逐台记录。"),
        ("收款与定金", "按实际收款方式拆分记录，可选择定金库中的可用定金抵扣。"),
        ("国补/教育补贴", "支持补贴金额、人员信息、订单号、凭证照片和开票资料。"),
        ("订单操作", "支持一键清空、缓存恢复和提交；提交后可按需打印小票。"),
    ], [2200, 7160])
    heading(d, "3.2 商品编码识别", 2)
    para(d, "支持拍照或从相册选择图片，通过 OCR 识别 PN、MTM、SN 等编码；识别结果可人工修正，也可切换到手动输入。")
    heading(d, "3.3 订单查询与详情", 2)
    para(d, "支持按订单编号、提交人、PN、SN、开票信息、归档状态、资源类型和日期范围筛选，支持查询全部、重置条件、导出订单和查看导出任务。详情页可按权限修改；归档后通常只开放国补人信息、国补照片和开票信息等限定字段。")
    heading(d, "3.4 定金管理", 2)
    para(d, "登记客户定金并形成定金库，支持按状态和客户电话/会员 ID 查询，符合条件的定金可办理退款登记，也可在订单中抵扣。")
    heading(d, "四、库存与供应链")
    heading(d, "4.1 库存查询", 2)
    para(d, "按 PN、SN 或商品名称查询当前门店、其他门店或全部门店库存，并查看定价、近 7 天/30 天销售等信息。")
    heading(d, "4.2 采购入库", 2)
    para(d, "查看待入库单，在详情中录入 PN、SN、库存类型、入库库位和数量。SN 商品逐台录入；无 SN 商品的分配数量应等于待入库数量。支持保存草稿后确认入库。")
    heading(d, "4.3 调拨管理", 2)
    para(d, "支持门店间申请调拨、选择实际出库库存、上传出库凭证、确认收货入库，并提供拒绝、退回和撤销等异常处理。调拨门店通常需属于同一经销商和区域。")
    heading(d, "五、审批与经营分析")
    heading(d, "5.1 审批中心", 2)
    para(d, "按角色和门店权限汇总待办，支持销售、采购、报销、商品、退库、销售退单、资源套回和毛利调整等业务类型，并提供历史记录查询。")
    heading(d, "5.2 经营报表", 2)
    para(d, "提供销售金额、成交订单、总毛利、毛利率、客单价、库存金额、门店排行、品类销售、重点产品和员工销售等分析，并支持今日、本周、本月或自定义日期范围。页面标注数据来源为 ANY-ERP 统计表。")
    heading(d, "5.3 业绩上报与商场销量", 2)
    para(d, "业绩上报统计查看订单已上报/未上报、数量、金额和归档时间；商场销量查询支持按订单、提交人、PN/SN、开票信息、资源类型、归档状态和日期查询，并支持导出和“我的任务”。")
    heading(d, "六、采购、费用与基础管理")
    table(d, ["模块", "主要能力"], [
        ("采购申请", "供应商、付款方式、发票类型、货型、采购商品、收货门店、库位、供应商截图和审批流程。"),
        ("费用管理", "费用发生方、金额、日期、支付方式、发票信息、费用说明和个人费用记录。"),
        ("门店/经销商信息", "新增、编辑门店，维护店长/店员和经销商员工信息。"),
        ("商品管理", "按名称、PN、SN 查询，维护 PN、SN、MTM、价格和分类。"),
        ("批量上传", "解析 CSV/TXT，展示错误、 新增、更新和失败结果。"),
        ("业务字典", "维护客户来源、收款方式和金额补录项目，统一录入口径。"),
        ("用户管理", "查看角色和门店，重置密码或恢复默认密码。"),
    ], [2200, 7160])
    heading(d, "七、权限、设备与注意事项")
    para(d, "经销商通常拥有跨门店管理能力；店长以所属门店管理为主；店员以日常业务执行为主。实际权限以账号配置为准。打印机需要开启手机蓝牙、授权微信蓝牙权限、搜索并绑定热敏打印机。")
    callout(d, "业务重点", "提交订单前重点核对商品编码、客户与收款信息、补贴资料和凭证照片；使用报表时关注数据来源、同步时间、日期范围和账号可见范围。", GOLD, "A26A00")
    d.save(ROOT / "小艾2.0产品说明书.docx")


def build_user():
    d = Document(); setup(d, "用户操作手册"); cover(d, "门店一线用户操作手册", "适用对象：店员、店长及需要日常操作小艾的门店人员")
    heading(d, "使用前先看")
    para(d, "本手册按门店日常工作顺序编写。如果某个按钮或菜单没有显示，通常是当前账号角色、门店范围或业务状态不满足条件，请联系经销商管理员确认。")
    table(d, ["角色", "常见工作"], [("店员", "新建订单、商品识别、订单查询、库存查询、发起调拨、查看个人信息"), ("店长", "完成店员工作，并处理本店入库、调拨、审批和经营查看"), ("经销商", "管理门店、商品、用户、基础资料、审批和跨门店经营分析")], [1800, 7560])
    heading(d, "一、登录与进入首页")
    for x in [
        "打开微信小程序，进入“登录授权”页面。",
        "输入账号和密码；首次使用时按管理员提供的账号登录，默认密码以页面提示为准。",
        "登录成功后核对姓名、账号类型、手机号和当前门店。",
        "如果账号有多个门店，在“个人中心”使用“切换”选择当前工作的门店。",
    ]: number(d, x)
    callout(d, "注意", "开单、库存和调拨前一定要确认当前门店；门店切换错误会导致查不到数据或提交到错误门店。", GOLD, "A26A00")
    heading(d, "二、新建销售订单")
    heading(d, "2.1 进入开单页面", 2)
    number(d, "首页点击“新建订单”。")
    number(d, "先填写客户信息，再录入商品，最后填写金额和收款。")
    heading(d, "2.2 填写客户信息", 2)
    for x in ["选择一级客户来源；如有二级来源，再继续选择二级来源。", "填写会员称呼和会员 ID；没有会员信息时按门店业务要求填写。"]: bullet(d, x)
    heading(d, "2.3 添加商品", 2)
    for x in [
        "点击“添加商品”，使用商品名称、PN 或 SN 搜索；也可以点击“扫码”。",
        "如果商品编码在实物上，优先使用“商品编码识别”拍照或从相册识别。",
        "核对商品名称、PN/SN、单价和数量；有 IMEI 字段时按要求填写。",
        "一笔订单有多个商品时重复添加，删除错误商品后再继续。",
    ]: number(d, x)
    heading(d, "2.4 填写收款、定金和补贴", 2)
    for x in [
        "在金额汇总区确认订单总计和应收金额。",
        "按实际情况添加现金、移动支付、银行卡等收款方式，并核对各项金额合计。",
        "需要使用定金时，点击定金选择框，选择客户可用定金并填写本次抵扣金额。",
        "涉及国补或教育补贴时，填写补贴人、ID/订单号、补贴金额及相关照片；需要开票时填写开票信息和金额。",
    ]: number(d, x)
    heading(d, "2.5 提交订单", 2)
    for x in ["提交前逐项检查商品编码、金额、收款方式和凭证。", "点击提交，等待页面提示成功；不要在提交过程中重复点击。", "需要小票时，进入订单详情或按页面提示打印。"]: number(d, x)
    heading(d, "三、查询订单与查看详情")
    for x in [
        "首页进入“订单查询”。",
        "按需要填写订单编号、PN、SN、开票信息或提交人；也可以选择归档状态、资源类型和日期范围。",
        "点击“查询”；要扩大范围时使用“查询全部”，要重新开始时点击“重置”。",
        "点击订单卡片进入“订单详情”，查看客户、商品、金额、收款和照片。",
        "有修改权限时进入编辑模式保存；归档订单通常只能补充国补和开票等限定信息。",
    ]: number(d, x)
    heading(d, "四、定金操作")
    heading(d, "4.1 登记定金", 2)
    for x in ["首页点击“定金管理”。", "填写客户姓名、会员 ID/电话、定金金额、收款方式和备注。", "点击“提交并存入定金库”，确认页面提示成功。"]: number(d, x)
    heading(d, "4.2 查询或退款登记", 2)
    for x in ["在定金单列表选择状态，或输入客户电话/会员 ID 查询。", "核对定金单号、客户、金额、收款方式和状态。", "对页面显示可退款的定金，点击“退款登记”，按提示完成。"]: number(d, x)
    heading(d, "五、查询库存")
    for x in ["首页点击“库存查询”。", "输入 PN、SN 或商品名称，点击“查询”。", "在“当前门店/其他门店/全部门店”之间切换查看范围。", "查看库存数量、定价和近 7 天/30 天销售，再决定销售、调拨或采购。"]: number(d, x)
    heading(d, "六、发起库存调拨")
    for x in [
        "进入“调拨管理”，点击“发起调拨申请”。",
        "选择调出门店和接收门店；系统会校验门店范围。",
        "点击“选择调拨商品”，按商品名称、编码、PN 或 SN 搜索，确认商品编码和库存。",
        "填写申请数量；SN 商品按页面要求选择具体 SN。",
        "点击“提交调拨申请”，在调拨单列表查看状态。",
        "出库时点击“选择商品并出库”，选择实际 PN/SN、数量并上传出库凭证。",
        "收到货后点击“确认收货入库”，核对商品并上传收货凭证。",
    ]: number(d, x)
    callout(d, "调拨限制", "调拨通常只支持同一经销商、同一区域内的门店；如果门店无法选择，请联系管理员检查门店归属和区域配置。", GOLD, "A26A00")
    heading(d, "七、查看采购入库单")
    for x in [
        "进入“采购入库”，选择入库门店和单据状态。",
        "点击待入库单进入详情，核对来源单号、门店、数量和金额。",
        "按商品逐项录入厂商编码 PN；SN 商品逐台输入或扫码。",
        "选择库存类型和入库库位，填写入库数量；必要时点击“添加库存类型”。",
        "信息未确认时点击“保存草稿”；确认无误后点击“确认入库”。",
    ]: number(d, x)
    heading(d, "八、处理审批")
    for x in [
        "进入“审批中心”，先查看“待审批”数量。",
        "使用销售、采购、报销、商品、退单等类型筛选，点击单据进入详情。",
        "核对申请人、门店、金额、商品和照片等业务明细。",
        "同意时点击“审批通过”；不同意时点击“拒绝”并填写明确原因。",
        "需要回看历史时切换到“商品审批记录”，输入申请单号、商品名、PN 或申请人查询。",
    ]: number(d, x)
    heading(d, "九、蓝牙打印机设置")
    for x in [
        "手机打开蓝牙，并允许微信使用蓝牙。",
        "进入“个人中心”→“打印机设置”，点击开启蓝牙或搜索设备。",
        "确认热敏打印机已开机且处于可发现状态。",
        "在设备列表点击“绑定”；绑定成功后订单打印时可自动连接。",
        "搜索不到设备时，确认打印机没有被其他手机连接，并重新授权蓝牙权限。",
    ]: number(d, x)
    heading(d, "十、常见问题")
    table(d, ["问题", "处理办法"], [
        ("商品搜不到", "确认搜索类型；检查 PN/SN 是否正确；联系管理员确认商品是否已维护或采购审批是否已通过。"),
        ("扫码/OCR 不准", "重新拍摄，保持光线充足和手机稳定；也可手动输入 PN、MTM 或 SN。"),
        ("订单不能编辑", "检查订单是否已归档、是否本人提交、是否属于当前门店，以及当前账号是否有修改权限。"),
        ("看不到审批单", "审批单按角色和门店范围显示；先确认当前门店和账号角色。"),
        ("调拨无法提交", "检查调出/接收门店、商品库存、数量和经销商/区域范围。"),
        ("打印失败", "检查蓝牙、微信权限、打印机电量和绑定状态，必要时重新搜索绑定。"),
    ], [2500, 6860])
    callout(d, "操作原则", "先核对编码，再核对金额，最后提交流程；涉及库存、补贴、报销和调拨时，务必保留清晰的照片凭证。", LIGHT, GREEN)
    d.save(ROOT / "小艾2.0用户操作手册.docx")


if __name__ == "__main__":
    build_product()
    build_user()
    print("done")
